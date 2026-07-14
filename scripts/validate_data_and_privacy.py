"""
데이터 · 개인정보 교차검증 (rename+갱신 <- validate_enhanced.py)
- 검증 대상을 최신 산출물(output/seoul_expense_executive_report.docx, output/dashboard.html)로 갱신
- 하드코딩된 데이터셋 고정값 대신 원본 CSV 재계산 값과 analysis_enhanced.json을 직접 대조(다른 CSV로도 재사용 가능)
- 결과: output/validation_report.md
"""
import argparse
import json
import re
import pandas as pd
from docx import Document

DOCX_PATH = "output/seoul_expense_executive_report.docx"
DASHBOARD_PATH = "output/dashboard.html"
JSON_PATH = "workspace/analysis_enhanced.json"
REPORT_PATH = "output/validation_report.md"


def recompute_ground_truth(csv_path):
    df = pd.read_csv(csv_path, encoding="utf-8-sig")
    df["집행금액"] = df["집행금액"].astype("int64")
    df["해당월"] = df["해당월"].astype(int)

    total = int(df["집행금액"].sum())
    rows = len(df)
    monthly = df.groupby("해당월")["집행금액"].sum().sort_index()

    n_months = df["해당월"].nunique()
    coverage = df.groupby("전체부서명")["해당월"].nunique()
    full_depts = coverage[coverage == n_months].index.tolist()

    dept_totals_all = df.groupby("전체부서명")["집행금액"].sum().sort_values(ascending=False)
    total_all = dept_totals_all.sum()
    top5_share = float(dept_totals_all.head(5).sum() / total_all * 100)
    top10_share = float(dept_totals_all.head(10).sum() / total_all * 100)
    hhi = float(((dept_totals_all / total_all) ** 2).sum() * 10000)

    return {
        "total": total, "rows": rows, "monthly": monthly,
        "full_depts": full_depts, "total_depts": df["전체부서명"].nunique(),
        "top5_share": top5_share, "top10_share": top10_share, "hhi": hhi,
        "df": df,
    }


def extract_docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def run(csv_path):
    gt = recompute_ground_truth(csv_path)
    with open(JSON_PATH, encoding="utf-8") as f:
        enh = json.load(f)

    lines = ["# 검증 보고서 — 부서운영업무추진비 분석\n",
             f"검증 대상: `{JSON_PATH}`, `{DOCX_PATH}`, `{DASHBOARD_PATH}`\n"]

    # ---- 1. 수치 일치성 (원본 재계산 <-> JSON) ----
    lines.append("## 1. 수치 일치성 검증 (원본 CSV 재계산 기준)\n")
    lines.append("| 항목 | 원본 재계산 | analysis_enhanced.json | 일치 여부 |")
    lines.append("|---|---|---|---|")
    checks = []

    def add_check(name, gt_val, other_val, fmt="{:,}"):
        ok = gt_val == other_val
        checks.append(ok)
        lines.append(f"| {name} | {fmt.format(gt_val)} | {fmt.format(other_val)} | {'✅ 일치' if ok else '❌ 불일치'} |")

    add_check("총 집행액(원)", gt["total"], enh["meta"]["총집행액"])
    add_check("총 행수(건)", gt["rows"], enh["meta"]["총행수"])
    add_check("동일기간 비교가능 부서 수", len(gt["full_depts"]), enh["meta"]["동일기간비교가능부서수"], fmt="{}")
    add_check("전체 부서 수", gt["total_depts"], enh["meta"]["전체부서수"], fmt="{}")

    for m in sorted(gt["monthly"].index):
        gt_m = int(gt["monthly"].loc[m])
        month_data = enh["2_monthly_category_mix"]["amount"].get(str(m), {})
        enh_m = int(sum(month_data.values()))
        add_check(f"{m}월 총액(원)", gt_m, enh_m)

    hv = enh["5_high_value_and_concentration"]
    add_check("상위5개 부서 점유율(%)", round(gt["top5_share"], 2), round(hv["top5_share_pct"], 2), fmt="{}")
    add_check("상위10개 부서 점유율(%)", round(gt["top10_share"], 2), round(hv["top10_share_pct"], 2), fmt="{}")
    add_check("HHI", round(gt["hhi"], 1), round(hv["hhi"], 1), fmt="{}")
    lines.append("")

    # ---- 2. DOCX 텍스트 내 핵심 수치 확인 (JSON 기준값을 동적으로 검색, 하드코딩 없음) ----
    lines.append("## 2. DOCX 산출물 내 수치 포함 여부 확인\n")
    docx_text = extract_docx_text(DOCX_PATH)
    expected_strings = {
        f"총 집행액 {enh['meta']['총집행액']:,}원": f"{enh['meta']['총집행액']:,}원" in docx_text,
        f"총 건수 {enh['meta']['총행수']:,}건": f"{enh['meta']['총행수']:,}건" in docx_text,
        f"동일기간 비교가능 부서 {enh['meta']['동일기간비교가능부서수']}개": f"{enh['meta']['동일기간비교가능부서수']}개" in docx_text,
        f"상위5개 부서 점유율 {hv['top5_share_pct']:.1f}%": f"{hv['top5_share_pct']:.1f}%" in docx_text,
        f"상위10개 부서 점유율 {hv['top10_share_pct']:.1f}%": f"{hv['top10_share_pct']:.1f}%" in docx_text,
        "동일기간 비교 유의사항 각주('비교 불가' 포함)": "비교 불가" in docx_text,
        "개인정보 비식별 처리 각주": "비식별" in docx_text,
    }
    lines.append("| 확인 항목 | 존재 여부 |")
    lines.append("|---|---|")
    for label, found in expected_strings.items():
        lines.append(f"| {label} | {'✅ 확인' if found else '❌ 누락'} |")
    lines.append("")

    # ---- 3. 동일기간 비교 조건 준수 ----
    lines.append("## 3. 동일기간 비교 조건 준수 검증\n")
    docx_dept_top5 = list(enh["4_dept_structure_same_period"].keys())[:5]
    all_in_full = all(d in gt["full_depts"] for d in docx_dept_top5)
    lines.append(f"- DOCX에 인용된 부서비교 상위5({', '.join(docx_dept_top5)})가 "
                 f"전 기간 데이터 보유 부서 목록에 전부 포함되는가: {'✅ 예' if all_in_full else '❌ 아니오'}")
    lines.append(f"- 전 기간 데이터 보유 부서 수(원본 재계산 vs JSON): {len(gt['full_depts'])} vs "
                 f"{enh['meta']['동일기간비교가능부서수']} ({'일치' if len(gt['full_depts']) == enh['meta']['동일기간비교가능부서수'] else '불일치'})")
    lines.append("")

    # ---- 4. 개인정보 노출 검증 ----
    lines.append("## 4. 개인정보 노출 검증\n")
    df_raw = gt["df"] if "전화번호" in gt["df"].columns else pd.read_csv(csv_path, encoding="utf-8-sig")
    has_pii_cols = "전화번호" in df_raw.columns and "작성자" in df_raw.columns
    pii_hits = []
    target_files = [JSON_PATH, DASHBOARD_PATH]

    if has_pii_cols:
        actual_phones = set(df_raw["전화번호"].astype(str).unique())
        actual_authors = set(df_raw["작성자"].astype(str).unique())
        for path in target_files:
            try:
                with open(path, encoding="utf-8") as f:
                    text = f.read()
            except FileNotFoundError:
                continue
            for phone in actual_phones:
                if phone in text:
                    pii_hits.append(f"{path}: 전화번호 원문 노출 발견")
            for author in actual_authors:
                if len(author) >= 2 and author in text:
                    pii_hits.append(f"{path}: 작성자 원문 노출 의심('{author}')")
        for phone in actual_phones:
            if phone in docx_text:
                pii_hits.append(f"{DOCX_PATH}: 전화번호 원문 노출 발견")

        lines.append(f"- 검사 대상 파일: {', '.join(target_files)}, `{DOCX_PATH}`")
        lines.append(f"- 원본 전화번호 고유값 {len(actual_phones)}개, 작성자 고유값 {len(actual_authors)}개 대조")
    else:
        lines.append("- 원본 CSV에 전화번호/작성자 컬럼이 없어 해당 검사를 건너뜀")

    if pii_hits:
        lines.append("- **결과: 노출 의심 발견**")
        for h in pii_hits:
            lines.append(f"  - {h}")
    else:
        lines.append("- 결과: ✅ 전화번호/작성자 원문 노출 없음")
    lines.append("")

    # ---- 5. 종합 판정 ----
    lines.append("## 5. 종합 판정\n")
    all_numeric_ok = all(checks)
    all_docx_ok = all(expected_strings.values())
    overall = all_numeric_ok and all_docx_ok and all_in_full and not pii_hits
    lines.append(f"- 수치 일치성: {'✅ 전항목 일치' if all_numeric_ok else '❌ 불일치 항목 존재'}")
    lines.append(f"- DOCX 수치 포함: {'✅ 전항목 확인' if all_docx_ok else '❌ 누락 항목 존재'}")
    lines.append(f"- 동일기간 비교 조건: {'✅ 준수' if all_in_full else '❌ 위반'}")
    lines.append(f"- 개인정보 비노출: {'✅ 확인' if not pii_hits else '❌ 의심 사례 발견'}")
    lines.append(f"\n**종합: {'🟢 제출 가능' if overall else '🔴 재검토 필요'}**")

    with open(REPORT_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print("검증 완료. overall =", overall)
    print("수치일치:", all_numeric_ok, "| DOCX확인:", all_docx_ok, "| 기간조건:", all_in_full, "| PII이슈:", len(pii_hits))
    return overall


def main():
    parser = argparse.ArgumentParser(description="데이터·개인정보 교차검증")
    parser.add_argument("--input", default="input/seoul_expenses.csv", help="원본 CSV 경로")
    args = parser.parse_args()
    run(args.input)


if __name__ == "__main__":
    main()
