"""
DOCX 레이아웃 구조 검증 (XML 레벨)
- output/layout_validation_report.md

주의: 이 환경에는 MS Word/LibreOffice가 없어 실제 렌더링 페이지 수는 확인할 수 없다.
따라서 이 스크립트는 "구조적 설정값"(여백, 표 너비, cantSplit, tblHeader, 이미지 크기, 페이지구분 수)을
XML 레벨에서 점검하는 정적 검증이며, 실제 렌더링 결과와 다를 수 있다는 한계를 보고서에 명시한다.
"""
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = "output/seoul_expense_executive_report.docx"
CONTENT_WIDTH_CM = 18.0
DXA_PER_CM = 567


def check_margins(doc):
    s = doc.sections[0]
    mm = lambda v: round(v.mm, 1)
    ok = mm(s.top_margin) == 15.0 and mm(s.bottom_margin) == 15.0 and mm(s.left_margin) == 15.0 and mm(s.right_margin) == 15.0
    return ok, f"상/하/좌/우 = {mm(s.top_margin)}/{mm(s.bottom_margin)}/{mm(s.left_margin)}/{mm(s.right_margin)}mm"


def check_tables(doc):
    results = []
    for i, table in enumerate(doc.tables):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        width_cm = int(tblW.get(qn('w:w'))) / DXA_PER_CM if tblW is not None else None
        width_ok = width_cm is not None and width_cm <= CONTENT_WIDTH_CM + 0.05

        rows = table.rows
        cant_split_count = 0
        for row in rows:
            trPr = row._tr.find(qn('w:trPr'))
            if trPr is not None and trPr.find(qn('w:cantSplit')) is not None:
                cant_split_count += 1
        cant_split_ok = cant_split_count == len(rows)

        header_row = rows[0]
        trPr0 = header_row._tr.find(qn('w:trPr'))
        header_repeat_ok = trPr0 is not None and trPr0.find(qn('w:tblHeader')) is not None

        results.append({
            "index": i, "rows": len(rows), "cols": len(table.columns),
            "width_cm": round(width_cm, 2) if width_cm else None, "width_ok": width_ok,
            "cant_split_ok": cant_split_ok, "cant_split_count": cant_split_count,
            "header_repeat_ok": header_repeat_ok,
        })
    return results


def check_images(doc):
    results = []
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            pass
    # 인라인 이미지의 실제 표시 폭은 run 내 drawing extent에서 추출
    from docx.oxml.ns import qn as q
    body = doc.element.body
    for i, drawing in enumerate(body.iter(q('w:drawing'))):
        extent = drawing.find('.//' + q('wp:extent'))
        if extent is not None:
            cx = int(extent.get('cx'))
            width_cm = cx / 360000  # EMU -> cm
            results.append({"index": i, "width_cm": round(width_cm, 2),
                             "within_95pct": width_cm <= CONTENT_WIDTH_CM * 0.95 + 0.05})
    return results


def check_page_breaks(doc):
    count = 0
    from docx.oxml.ns import qn as q
    for br in doc.element.body.iter(q('w:br')):
        if br.get(q('w:type')) == 'page':
            count += 1
    return count


def check_fonts(doc):
    sizes = []
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.size:
                sizes.append(r.font.size.pt)
    # 7~7.5pt: 캡션/각주(의도적 축소), 8~9pt: 표, 9.5~10pt: 본문, 12~19pt: 제목류 — 전부 의도된 설계값
    return {
        "min": min(sizes) if sizes else None, "max": max(sizes) if sizes else None,
        "body_range_ok": all(7 <= s <= 20 for s in sizes),
    }


def run():
    doc = Document(DOCX_PATH)
    lines = ["# DOCX 레이아웃 검증 보고서\n",
             f"검증 대상: `{DOCX_PATH}`\n",
             "> **한계**: 이 환경에는 MS Word/LibreOffice가 설치되어 있지 않아 실제 렌더링 페이지 수·페이지별 배치를 "
             "기계적으로 확인할 수 없습니다. 아래는 DOCX XML에 실제로 기록된 설정값(여백, 표 너비, 행 분할 방지, "
             "헤더 반복, 이미지 크기, 페이지구분 삽입 수)에 대한 **정적 구조 검증**입니다.\n"]

    lines.append("## 1. 페이지 여백\n")
    ok, detail = check_margins(doc)
    lines.append(f"- 결과: {'✅ 15mm 통일 확인' if ok else '❌ 불일치'} ({detail})\n")

    lines.append("## 2. 표 검증 (열 너비 / 행 분할 방지 / 헤더 반복)\n")
    lines.append("| 표# | 행수 | 열수 | 표 너비(cm) | 본문폭(18cm) 이내 | 전 행 cantSplit | 헤더 tblHeader |")
    lines.append("|---|---|---|---|---|---|---|")
    tables = check_tables(doc)
    all_tables_ok = True
    for t in tables:
        w_ok = "✅" if t["width_ok"] else "❌"
        cs_ok = "✅" if t["cant_split_ok"] else f"❌ ({t['cant_split_count']}/{t['rows']})"
        hr_ok = "✅" if t["header_repeat_ok"] else "❌"
        if not (t["width_ok"] and t["cant_split_ok"] and t["header_repeat_ok"]):
            all_tables_ok = False
        lines.append(f"| {t['index']+1} | {t['rows']} | {t['cols']} | {t['width_cm']} | {w_ok} | {cs_ok} | {hr_ok} |")
    lines.append("")

    lines.append("## 3. 이미지(차트) 크기 검증 (본문폭 95% = 17.1cm 이내)\n")
    images = check_images(doc)
    lines.append("| 이미지# | 폭(cm) | 95% 이내 |")
    lines.append("|---|---|---|")
    all_images_ok = True
    for im in images:
        ok_mark = "✅" if im["within_95pct"] else "❌"
        if not im["within_95pct"]:
            all_images_ok = False
        lines.append(f"| {im['index']+1} | {im['width_cm']} | {ok_mark} |")
    lines.append("")

    lines.append("## 4. 페이지구분 삽입 수\n")
    pb = check_page_breaks(doc)
    lines.append(f"- 명시적 페이지구분(`w:br type=page`) {pb}회 삽입 → 최소 {pb+1}개 섹션 블록으로 구성됨을 의미\n"
                 f"  (표·차트 자체 분량에 따라 실제 렌더링 페이지 수는 이보다 많을 수 있음 — 위 한계 참고)\n")

    lines.append("## 5. 폰트 크기 범위\n")
    fonts = check_fonts(doc)
    lines.append(f"- 사용된 폰트 크기 범위: {fonts['min']}pt ~ {fonts['max']}pt "
                 f"({'✅ 7~20pt 범위 내(7~8pt는 캡션·각주 의도적 축소, 9.5~10pt 본문, 8~9pt 표)' if fonts['body_range_ok'] else '❌ 범위 이탈'})\n")

    lines.append("## 6. 종합 판정\n")
    overall = all_tables_ok and all_images_ok and ok
    lines.append(f"- 여백: {'✅' if ok else '❌'} | 표 규칙: {'✅' if all_tables_ok else '❌'} | "
                 f"이미지 크기: {'✅' if all_images_ok else '❌'}")
    lines.append(f"\n**구조적 검증 종합: {'🟢 규칙 준수 확인' if overall else '🔴 일부 규칙 위반'}**")
    lines.append("\n> 실제 페이지 넘김 시 표/차트가 시각적으로 잘리는지는 Word 등에서 직접 열어 육안으로 최종 확인이 필요합니다.")

    with open("output/layout_validation_report.md", "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print("검증 완료. overall =", overall, "| 표:", len(tables), "| 이미지:", len(images), "| 페이지구분:", pb)
    return overall


def main():
    run()


if __name__ == "__main__":
    main()
