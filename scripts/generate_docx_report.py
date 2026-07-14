"""
DOCX 상세 보고서 생성 (rename+데이터구동화 <- generate_docx_detailed_v2.py)
- output/seoul_expense_executive_report.docx

기존 generate_docx_detailed_v2.py 대비 변경점:
- 본문에 하드코딩되어 있던 수치·서술을 전부 workspace/analysis_enhanced.json에서 동적으로 도출하도록 변경
  (다른 CSV로 재실행해도 보고서 내용이 실제 데이터를 반영하도록)
- 레이아웃 규칙(표 열너비 명시, 행 분할 방지, 헤더 반복, keep_with_next)은 동일하게 유지

개인정보(전화번호/작성자/집행대상 원문) 미포함. 고액 집행은 오류로 단정하지 않음.
"""
import argparse
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

JSON_PATH = "workspace/analysis_enhanced.json"
OUT_PATH = "output/seoul_expense_executive_report.docx"
CH = "output/charts"
FONT = "맑은 고딕"
CONTENT_WIDTH_CM = 21.0 - 1.5 - 1.5  # A4 폭 - 좌우 여백 15mm*2 = 18.0cm


# ---------------- 공통 함수 (레이아웃 규칙, 데이터셋 무관) ----------------

def set_page_layout(doc):
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    return section


def set_font(run, size=10, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT)


def add_p(doc, text="", size=10, bold=False, align=None, color=None,
          space_after=3, space_before=0, keep_with_next=False, keep_together=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align:
        p.alignment = align
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    if keep_together:
        p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_font(run, size, bold, color)
    return p


def add_safe_page_break(doc):
    doc.add_page_break()


def add_spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    p.paragraph_format.space_before = Pt(0)


def set_table_width(table, total_width_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int(total_width_cm * 567)))


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))


def set_col_widths(table, widths_cm):
    set_table_width(table, sum(widths_cm))
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            set_cell_width(cell, w)


def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn('w:tblHeader')) is None:
        trPr.append(OxmlElement('w:tblHeader'))


def apply_no_split_to_all_rows(table, header_rows=1):
    for i, row in enumerate(table.rows):
        prevent_row_split(row)
        if i < header_rows:
            set_repeat_table_header(row)


def add_cell_text(cell, text, size=9, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def format_currency_cell(cell, amount, size=8.5):
    add_cell_text(cell, f"{int(amount):,}원", size=size, align=WD_ALIGN_PARAGRAPH.RIGHT)


def add_chart_with_caption(doc, path, width_cm, caption):
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.keep_with_next = True
    img_p.add_run().add_picture(path, width=Cm(width_cm))
    add_p(doc, caption, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100), space_after=6)


def section_title(doc, text, num=None):
    label = f"{num}. {text}" if num else text
    return add_p(doc, label, size=13.5, bold=True, space_after=5, space_before=2, keep_with_next=True)


def format_department_table(doc, dept_items):
    widths = [1.2, 6.0, 3.5, 1.8, 3.0, 2.5]
    table = doc.add_table(rows=1 + len(dept_items), cols=6)
    table.style = "Light Grid Accent 1"
    set_col_widths(table, widths)
    hdr = ["순위", "전체부서명", "집행액", "건수", "최대비목(비중)", "변동계수"]
    for j, h in enumerate(hdr):
        add_cell_text(table.cell(0, j), h, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, (name, v) in enumerate(dept_items, start=1):
        add_cell_text(table.cell(i, 0), str(i), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(table.cell(i, 1), name, size=8)
        format_currency_cell(table.cell(i, 2), v["총액"])
        add_cell_text(table.cell(i, 3), f"{int(v['건수'])}건", size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_cell_text(table.cell(i, 4), f"{v.get('최대비목','-')}({v.get('최대비목비중(%)',0):.1f}%)", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(table.cell(i, 5), f"{v.get('변동계수',0):.2f}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_no_split_to_all_rows(table, header_rows=1)
    return table


def format_category_table(doc, cat_rows):
    widths = [4.0, 5.0, 3.0]
    table = doc.add_table(rows=1 + len(cat_rows), cols=3)
    table.style = "Light Grid Accent 1"
    set_col_widths(table, widths)
    for j, h in enumerate(["비목", "집행액", "비중"]):
        add_cell_text(table.cell(0, j), h, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, (a, b, c) in enumerate(cat_rows, start=1):
        add_cell_text(table.cell(i, 0), a, size=8.5)
        add_cell_text(table.cell(i, 1), b, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_cell_text(table.cell(i, 2), c, size=8.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
    apply_no_split_to_all_rows(table, header_rows=1)
    return table


def format_kpi_table(doc, labels, values):
    n = len(labels)
    w = CONTENT_WIDTH_CM / n
    table = doc.add_table(rows=2, cols=n)
    set_col_widths(table, [w] * n)
    for i, (lab, val) in enumerate(zip(labels, values)):
        add_cell_text(table.cell(0, i), lab, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        add_cell_text(table.cell(1, i), val, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    apply_no_split_to_all_rows(table, header_rows=1)
    return table


# ---------------- 데이터 도출 (JSON -> 서술 재료, 데이터셋 무관) ----------------

def won(n):
    return f"{int(round(n)):,}원"


def derive_facts(enh):
    meta = enh["meta"]
    dept_items = list(enh["4_dept_structure_same_period"].items())

    # 비목별 표 (금액 내림차순)
    cat_amounts = {}
    for month_data in enh["2_monthly_category_mix"]["amount"].values():
        for cat, amt in month_data.items():
            cat_amounts[cat] = cat_amounts.get(cat, 0) + amt
    total = meta["총집행액"]
    cat_rows = sorted(cat_amounts.items(), key=lambda kv: -kv[1])
    cat_table_rows = [(c, won(a), f"{a/total*100:.1f}%") for c, a in cat_rows]

    # 월별 총액 -> 최고월
    monthly_totals = {m: sum(v.values()) for m, v in enh["2_monthly_category_mix"]["amount"].items()}
    top_month = max(monthly_totals, key=monthly_totals.get)

    # 전월대비 증감(문장)
    mom_items = sorted(enh["3_mom_change"].items(), key=lambda kv: int(kv[0]))
    mom_sentence_parts = []
    for m, v in mom_items:
        if v["전월대비증감률(%)"] is None:
            mom_sentence_parts.append(f"{m}월은 비교 불가")
        else:
            mom_sentence_parts.append(f"{m}월 {v['전월대비증감률(%)']:+.1f}%")
    mom_sentence = ", ".join(mom_sentence_parts)

    # 최대 증가월
    valid_mom = [(m, v["전월대비증감률(%)"]) for m, v in mom_items if v["전월대비증감률(%)"] is not None]
    max_mom_month, max_mom_val = max(valid_mom, key=lambda t: t[1]) if valid_mom else (None, None)

    # 비목 의존도 개수
    dep_threshold = 60
    dep_count = sum(1 for _, v in dept_items if v.get("최대비목비중(%)", 0) >= dep_threshold)
    full_dep_depts = [name for name, v in dept_items if v.get("최대비목비중(%)", 0) >= 99.5]

    # 변동계수 최고/최저
    cv_items = [(name, v.get("변동계수")) for name, v in dept_items if v.get("변동계수") is not None]
    max_cv = max(cv_items, key=lambda t: t[1]) if cv_items else None
    min_cv = min(cv_items, key=lambda t: t[1]) if cv_items else None

    # 집중도
    hv = enh["5_high_value_and_concentration"]
    top5_share = hv["top5_share_pct"]
    top10_share = hv["top10_share_pct"]
    hhi = hv["hhi"]

    # 고액 후보
    iqr_count = hv["iqr_candidate_count"]
    p95_count = hv["p95_candidate_count"]
    iqr_by_cat = hv["iqr_candidates_by_category"]
    top_iqr_cat = max(iqr_by_cat, key=iqr_by_cat.get) if iqr_by_cat else None
    top_iqr_cat_pct = (iqr_by_cat[top_iqr_cat] / iqr_count * 100) if iqr_count else 0
    iqr_by_month = hv["iqr_candidates_by_month"]
    top_iqr_month = max(iqr_by_month, key=iqr_by_month.get) if iqr_by_month else None
    top_iqr_month_pct = (iqr_by_month[top_iqr_month] / iqr_count * 100) if iqr_count and top_iqr_month else 0

    # 반복 패턴 1위
    repeats = enh["1_repeat_patterns"]
    top_repeat = repeats[0] if repeats else None

    return {
        "meta": meta, "dept_items": dept_items, "cat_table_rows": cat_table_rows,
        "top_month": top_month, "top_month_amount": monthly_totals[top_month],
        "mom_sentence": mom_sentence, "max_mom_month": max_mom_month, "max_mom_val": max_mom_val,
        "dep_count": dep_count, "dep_total": len(dept_items), "full_dep_depts": full_dep_depts,
        "max_cv": max_cv, "min_cv": min_cv,
        "top5_share": top5_share, "top10_share": top10_share, "hhi": hhi,
        "iqr_count": iqr_count, "p95_count": p95_count, "total_rows": meta["총행수"],
        "top_iqr_cat": top_iqr_cat, "top_iqr_cat_pct": top_iqr_cat_pct,
        "top_iqr_month": top_iqr_month, "top_iqr_month_pct": top_iqr_month_pct,
        "top_repeat": top_repeat,
    }


# ---------------- 본문 구성 ----------------

def build(enh):
    f = derive_facts(enh)
    meta = f["meta"]

    doc = Document()
    set_page_layout(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)

    # ===== 1페이지: 표지 / 개요 / KPI / 핵심 결과 =====
    add_p(doc, "부서운영업무추진비 분석 결과 (상세판)", size=19, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_p(doc, f"상세 분석 자료 | 분석기간 {meta['분석기간']}",
          size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=(90, 90, 90), space_after=12)

    section_title(doc, "분석 개요", 1)
    overview = [
        f"분석 대상 기간: {meta['분석기간']}",
        f"분석 대상 부서: 전체부서명 기준 {meta['전체부서수']}개 "
        f"(이 중 전 기간 데이터 보유 부서 {meta['동일기간비교가능부서수']}개만 상호 비교 가능)",
        f"분석 건수: 라인아이템 {meta['총행수']:,}건",
        "분석 기준: 해당년도→해당월→전체부서명→비목→집행금액 합계. 부서 간 비교는 동일 기간 데이터 보유 부서 간에만 수행",
    ]
    for o in overview:
        add_p(doc, f"· {o}", size=9.5, space_after=2)

    add_spacer(doc, 4)
    section_title(doc, "핵심 지표(KPI)", 2)
    format_kpi_table(doc,
        ["분석 기간", "총 집행액", "총 집행 건수", "비교가능 부서 수", "최고 집행월"],
        [meta["분석기간"], won(meta["총집행액"]), f"{meta['총행수']:,}건",
         f"{meta['동일기간비교가능부서수']}개(전체 {meta['전체부서수']}개 중)",
         f"{f['top_month']}월({won(f['top_month_amount'])})"])

    add_spacer(doc, 6)
    section_title(doc, "핵심 결과 요약", 3)
    highlights = []
    if f["max_mom_month"]:
        highlights.append(f"{f['max_mom_month']}월 집행액이 전월 대비 {f['max_mom_val']:+.1f}%로 가장 크게 변동")
    highlights.append(f"동일기간 비교 가능 {f['dep_total']}개 부서 중 {f['dep_count']}개 부서가 특정 비목에 60% 이상 의존")
    highlights.append(f"상위 5개 부서가 전체 집행액의 {f['top5_share']:.1f}%, 상위 10개 부서가 {f['top10_share']:.1f}%를 차지(HHI {f['hhi']:.1f})")
    for h in highlights:
        add_p(doc, f"· {h}", size=9.5, space_after=2)

    add_safe_page_break(doc)

    # ===== 2페이지: 월별 집행 현황 =====
    section_title(doc, "월별 집행 현황", 4)
    add_chart_with_caption(doc, f"{CH}/monthly_trend.png", 17.0,
        f"월별·부서별 집행 현황 (상위5부서+기타) | 분석기간 {meta['분석기간']} | 단위: 백만원 | 집계 기준: 해당월×전체부서명×집행금액 합계")
    add_p(doc, f"전월 대비 증감률: {f['mom_sentence']}", size=9, color=(80, 80, 80), space_after=4)

    add_safe_page_break(doc)

    # ===== 3페이지: 월별 비목 구성 변화 + 비목별 표 =====
    section_title(doc, "월별 비목 구성 변화", 5)
    add_chart_with_caption(doc, f"{CH}/monthly_category_mix_enhanced.png", 15.0,
        f"월별 비목 구성 변화(100% 환산) | 분석기간 {meta['분석기간']} | 단위: 비중(%) | 집계 기준: 해당월×비목×집행금액 비중")

    add_spacer(doc, 4)
    section_title(doc, "비목별 집행 현황", 6)
    format_category_table(doc, f["cat_table_rows"])

    add_safe_page_break(doc)

    # ===== 4페이지: 동일기간 부서 비교표 =====
    section_title(doc, f"동일기간({meta['분석기간']} 전기간 보유) 부서 비교", 7)
    format_department_table(doc, f["dept_items"])
    add_spacer(doc, 4)
    section_title(doc, "비교 조건과 해석", 8)
    interp = [
        f"전체 {meta['전체부서수']}개 부서 중 전 기간 데이터를 모두 보유한 {meta['동일기간비교가능부서수']}개 부서만 비교 대상으로 선정했다. "
        f"나머지 부서는 기간이 상이해 비교 불가로 처리한다.",
    ]
    if f["full_dep_depts"]:
        interp.append(f"{', '.join(f['full_dep_depts'])}는 단일 비목 비중이 100%로, 특정 비목에 대한 의존도가 가장 높다.")
    if f["max_cv"] and f["min_cv"]:
        interp.append(f"변동계수(월간표준편차/월간평균)는 {f['max_cv'][0]}({f['max_cv'][1]:.2f})이 최고, "
                       f"{f['min_cv'][0]}({f['min_cv'][1]:.2f})이 최저로, 월별 집행 규모의 상대적 변동 폭에 차이가 있다.")
    for t in interp:
        add_p(doc, f"· {t}", size=9.5, space_after=3)

    add_safe_page_break(doc)

    # ===== 5페이지: 부서×비목 히트맵 =====
    section_title(doc, "부서 × 비목 구성 히트맵", 9)
    add_chart_with_caption(doc, f"{CH}/dept_structure_heatmap_enhanced.png", 16.5,
        f"동일기간 부서 × 비목 구성비 | 분석기간 {meta['분석기간']} | 단위: 비중(%) | 집계 기준: 전체부서명×비목×집행금액 비중")

    add_safe_page_break(doc)

    # ===== 6페이지: 부서별 산점도 =====
    section_title(doc, "부서별 집행액 · 건수 산점도", 10)
    add_chart_with_caption(doc, f"{CH}/dept_scatter_enhanced.png", 16.5,
        f"동일기간 부서 규모-활동성 (점크기=건당평균, 색=주비목) | 분석기간 {meta['분석기간']} | 집계 기준: 전체부서명×집행금액·건수")

    add_safe_page_break(doc)

    # ===== 7페이지: 집중도 파레토 =====
    section_title(doc, "부서 집중도 파레토", 11)
    add_chart_with_caption(doc, f"{CH}/concentration_pareto_enhanced.png", 17.0,
        f"부서 집중도 파레토(전체 {meta['전체부서수']}개 부서 기준) | 분석기간 {meta['분석기간']} | 집계 기준: 전체부서명×집행금액 합계 및 누적 점유율")
    add_p(doc, f"집중도 지표(전체 {meta['전체부서수']}개 부서 기준): 상위5개 부서 점유율 {f['top5_share']:.1f}%, "
          f"상위10개 부서 점유율 {f['top10_share']:.1f}%, HHI {f['hhi']:.1f} — "
          "지출 분산·집중 정도를 나타내는 기술 지표이며, 값이 높다는 사실 자체를 부정적으로 해석하지 않는다. "
          "위 부서비교표와는 대상 범위가 다른 별개 지표다.", size=9, color=(80, 80, 80), space_after=4)

    add_safe_page_break(doc)

    # ===== 8페이지: 금액 분포 및 고액 집행 후보 =====
    section_title(doc, "금액 분포 및 고액 집행 후보", 12)
    add_chart_with_caption(doc, f"{CH}/amount_distribution_enhanced.png", 17.0,
        f"전체 집행금액 분포 및 고액 집행 후보 기준선 | 분석기간 {meta['분석기간']} | 단위: 원 | 집계 기준: 집행금액(전체 {f['total_rows']:,}건)")
    iqr_pct = f["iqr_count"] / f["total_rows"] * 100 if f["total_rows"] else 0
    p95_pct = f["p95_count"] / f["total_rows"] * 100 if f["total_rows"] else 0
    detail = (f"IQR 기준 고액 집행 후보 {f['iqr_count']}건({iqr_pct:.1f}%), 상위5% 기준 {f['p95_count']}건({p95_pct:.2f}%) "
              "— 두 기준을 구분해 제시하며, 오류나 부적정 집행으로 단정하지 않는다.")
    if f["top_iqr_month"]:
        detail += f" {f['top_iqr_month']}월에 {f['top_iqr_month_pct']:.1f}%로 집중"
    if f["top_iqr_cat"]:
        detail += f", 비목별로는 {f['top_iqr_cat']} 비목에 {f['top_iqr_cat_pct']:.1f}% 집중."
    add_p(doc, detail, size=9, color=(80, 80, 80), space_after=4)

    add_safe_page_break(doc)

    # ===== 9페이지: 반복 집행 패턴 =====
    section_title(doc, "반복 집행 패턴", 13)
    add_chart_with_caption(doc, f"{CH}/repeat_patterns_enhanced.png", 17.0,
        f"반복 집행 패턴 상위 10 (동일 부서·비목·금액 조합) | 분석기간 {meta['분석기간']} | 집계 기준: 전체부서명×비목×집행금액 반복횟수")
    repeat_note = "반복 사유나 정례성은 데이터로 확인되지 않아 단정하지 않는다."
    if f["top_repeat"]:
        tr = f["top_repeat"]
        repeat_note = (f"{tr['전체부서명']}의 {won(tr['집행금액'])}({tr['비목']}) 조합이 {tr['반복월수']}개월에 걸쳐 "
                       f"{tr['반복횟수']}회로 가장 많이 반복 관측됨. " + repeat_note)
    add_p(doc, repeat_note, size=9, color=(80, 80, 80), space_after=4)

    add_safe_page_break(doc)

    # ===== 10페이지: 핵심 시사점 / 분석 한계 =====
    section_title(doc, "핵심 시사점", 14)
    insights = list(highlights)
    if f["top_repeat"]:
        tr = f["top_repeat"]
        insights.append(f"{tr['전체부서명']}의 {won(tr['집행금액'])}({tr['비목']}) 조합은 {tr['반복월수']}개월에 걸쳐 {tr['반복횟수']}회 반복 관측됨")
    if f["top_iqr_cat"]:
        insights.append(f"고액 집행 후보(IQR 기준 {f['iqr_count']}건)의 {f['top_iqr_cat_pct']:.1f}%가 {f['top_iqr_cat']} 비목에서 발생")
    for ins in insights:
        add_p(doc, f"· {ins}", size=9.5, space_after=3)

    add_spacer(doc, 4)
    section_title(doc, "검토 사항 (데이터 품질 및 분석 한계)", 15)
    limits = [
        f"전체 {meta['전체부서수']}개 부서 중 {meta['동일기간비교가능부서수']}개 부서만 전 기간 데이터를 보유해 부서 간 순위 비교는 이 부서들로 한정됨",
        "데이터 기간이 단일 구간이라 전기 대비 비교 및 계절성 판단에는 한계가 있음",
        "집행장소·집행대상 등 일부 필드의 결측 건은 '데이터 확인 필요'로 처리, 금액 집계에는 영향 없음",
        "집행대상 컬럼은 업체명으로 추정되나 완전히 확정할 수 없어 원문을 노출하지 않음",
    ]
    for lim in limits:
        add_p(doc, f"· {lim}", size=9.5, space_after=3)

    add_spacer(doc, 6)
    footnote = ("집계 기준: 해당년도→해당월→전체부서명→비목→집행금액 합계 | 원본 데이터는 수정하지 않음 | "
                "개인정보 처리: 전화번호·작성자·집행대상 원문은 본 보고서에 포함하지 않음(비식별 처리) | "
                "고액 집행 후보는 오류·부적정 집행으로 단정하지 않음 | "
                "상세 인터랙티브 자료는 output/dashboard.html 참조")
    add_p(doc, footnote, size=7.5, color=(120, 120, 120), space_after=0)

    return doc


def run():
    with open(JSON_PATH, encoding="utf-8") as fp:
        enh = json.load(fp)
    doc = build(enh)
    doc.save(OUT_PATH)
    print(f"[generate_docx_report] 저장 완료: {OUT_PATH}")
    return OUT_PATH


def main():
    parser = argparse.ArgumentParser(description="DOCX 상세 보고서 생성 (workspace/analysis_enhanced.json 기반)")
    parser.parse_args()
    run()


if __name__ == "__main__":
    main()
